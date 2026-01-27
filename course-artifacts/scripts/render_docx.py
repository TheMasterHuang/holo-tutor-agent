from __future__ import annotations

from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from visual_spec import get_content_md, get_meta_date, get_meta_title, get_meta_watermark


def add_header_watermark(doc: Document, watermark: str) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(watermark)
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(200, 200, 200)
    run.font.bold = True


def _add_md_block(doc: Document, md: str) -> None:
    if not md:
        return
    text = md.replace("\r\n", "\n")
    in_code = False
    for raw in text.split("\n"):
        line = raw.rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            continue

        if in_code:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(10)
            continue

        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            continue

        if stripped == "":
            doc.add_paragraph("")
            continue

        doc.add_paragraph(line)


def render_lecture_docx(course_data: Dict[str, Any], out_docx_path: str) -> None:
    title = get_meta_title(course_data)
    date = get_meta_date(course_data)
    watermark = get_meta_watermark(course_data)

    sections: List[Dict[str, Any]] = course_data.get("lecture_notes") or course_data.get("sections") or []

    doc = Document()
    add_header_watermark(doc, watermark)

    doc.add_heading(f"{title} 讲稿", level=0)
    doc.add_paragraph(f"Date: {date}")
    doc.add_paragraph(f"Watermark: {watermark}")

    doc.add_heading("目录", level=1)
    for i, sec in enumerate(sections, start=1):
        doc.add_paragraph(f"{i}. {sec.get('title', '')}")

    for sec in sections:
        doc.add_page_break()
        doc.add_heading(sec.get("title", ""), level=1)
        _add_md_block(doc, get_content_md(sec))

    doc.save(out_docx_path)


def _normalize_true_false_answer(v: Any) -> str:
    if isinstance(v, bool):
        return "正确" if v else "错误"
    s = str(v).strip()
    if s in {"true", "True", "T", "正确"}:
        return "正确"
    if s in {"false", "False", "F", "错误"}:
        return "错误"
    return s


def _add_quiz_section(doc: Document, title: str, questions: List[Dict[str, Any]], *, kind: str) -> None:
    doc.add_heading(title, level=1)

    letters = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"
    for i, q in enumerate(questions, start=1):
        stem = (q.get("stem") or "").strip()
        doc.add_paragraph(f"{i}. {stem}")

        if kind == "single_choice":
            options = q.get("options") or []
            if isinstance(options, list):
                for j, opt in enumerate(options):
                    label = letters[j] if j < len(letters) else str(j + 1)
                    doc.add_paragraph(f"{label}. {str(opt).strip()}")

        if kind == "true_false":
            answer = _normalize_true_false_answer(q.get("answer"))
        else:
            answer = str(q.get("answer") or "").strip()

        explanation = str(q.get("explanation") or "").strip()
        doc.add_paragraph(f"答案：{answer}")
        doc.add_paragraph(f"解析：{explanation}")


def render_quiz_docx(course_data: Dict[str, Any], out_docx_path: str) -> None:
    title = get_meta_title(course_data)
    date = get_meta_date(course_data)
    watermark = get_meta_watermark(course_data)

    qb = course_data.get("quiz_bank") or {}
    single_choice = qb.get("single_choice") or []
    fill_blank = qb.get("fill_blank") or []
    true_false = qb.get("true_false") or []

    doc = Document()
    add_header_watermark(doc, watermark)

    doc.add_heading(f"{title} 习题集", level=0)
    doc.add_paragraph(f"Date: {date}")
    doc.add_paragraph(f"Watermark: {watermark}")

    if isinstance(single_choice, list):
        _add_quiz_section(doc, "一、单选题（10题）", single_choice, kind="single_choice")
    if isinstance(fill_blank, list):
        _add_quiz_section(doc, "二、填空题（10题）", fill_blank, kind="fill_blank")
    if isinstance(true_false, list):
        _add_quiz_section(doc, "三、判断题（10题）", true_false, kind="true_false")

    doc.save(out_docx_path)
